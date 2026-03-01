#!/usr/bin/env python3
"""
Docx to Markdown Converter
将 Word .docx 文件转换为 Markdown 格式
"""

import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import argparse


def docx_to_markdown(docx_path, output_path=None):
    """将docx文件转换为markdown格式"""
    
    doc = Document(docx_path)
    
    # 构建输出路径
    if output_path is None:
        base, _ = os.path.splitext(docx_path)
        output_path = base + ".md"
    
    markdown_content = []
    
    for paragraph in paragraph_iter(doc):
        # 获取段落文本
        text = get_paragraph_text(paragraph)
        if not text.strip():
            markdown_content.append("")
            continue
        
        # 获取段落格式
        heading_level = get_heading_level(paragraph)
        
        if heading_level:
            # 标题
            markdown_content.append(f"{'#' * heading_level} {text}")
        elif is_table(paragraph):
            # 表格
            continue  # 表格单独处理
        else:
            # 普通段落
            # 检查是否是列表
            if text.strip().startswith(('- ', '* ', '+ '):
                markdown_content.append(text)
            elif re.match(r'^\d+\.\s', text):
                markdown_content.append(text)
            else:
                # 处理加粗、斜体等
                formatted = format_text(paragraph)
                markdown_content.append(formatted)
    
    # 处理表格
    tables = [table for table in doc.tables]
    if tables:
        markdown_content.append("\n## 表格\n")
        for i, table in enumerate(tables):
            md_table = convert_table(table)
            markdown_content.append(f"\n### 表 {i+1}\n")
            markdown_content.append(md_table)
    
    # 写入文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_content))
    
    print(f"转换完成: {output_path}")
    return output_path


def paragraph_iter(doc):
    """迭代文档中的所有段落"""
    for element in doc.element.body:
        if element.tag.endswith('p'):  # paragraph
            yield paragraph_from_element(doc, element)


def paragraph_from_element(doc, element):
    """从element创建段落对象"""
    p = doc.add_paragraph()
    p._element.getparent().remove(p._element)
    p._element = element
    return p


def get_paragraph_text(paragraph):
    """获取段落文本"""
    texts = []
    for run in paragraph.runs:
        texts.append(run.text)
    return ''.join(texts)


def get_heading_level(paragraph):
    """获取标题级别"""
    try:
        style = paragraph.style.name if paragraph.style else ""
        if 'Heading' in style:
            # 尝试从样式名中提取级别
            match = re.search(r'Heading\s*(\d+)', style)
            if match:
                return int(match.group(1))
    except:
        pass
    return None


def is_table(paragraph):
    """检查是否是表格"""
    return False  # 简单实现


def convert_table(table):
    """转换表格为markdown格式"""
    rows = []
    
    # 表头
    header_row = []
    for cell in table.rows[0].cells:
        header_row.append(cell.text.strip() or " ")
    rows.append("| " + " | ".join(header_row) + " |")
    
    # 分隔行
    rows.append("| " + " | ".join(["---"] * len(header_row)) + " |")
    
    # 数据行
    for row in table.rows[1:]:
        data_row = []
        for cell in row.cells:
            data_row.append(cell.text.strip() or " ")
        rows.append("| " + " | ".join(data_row) + " |")
    
    return '\n'.join(rows)


def format_text(paragraph):
    """格式化文本，保留加粗、斜体等"""
    parts = []
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        
        # 检查格式
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        parts.append(text)
    
    result = ''.join(parts)
    
    # 检查段落对齐
    if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        result = f"\n<div align='center'>{result}</div>\n"
    
    return result


def main():
    parser = argparse.ArgumentParser(description='将Word docx文件转换为Markdown格式')
    parser.add_argument('input', help='输入的docx文件路径')
    parser.add_argument('-o', '--output', help='输出的md文件路径（可选）')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 - {args.input}")
        sys.exit(1)
    
    if not args.input.endswith('.docx'):
        print("错误: 请输入.docx格式的文件")
        sys.exit(1)
    
    docx_to_markdown(args.input, args.output)


if __name__ == '__main__':
    main()
